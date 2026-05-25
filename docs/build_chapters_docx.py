from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "chapters_3_to_5.md"
OUTPUT = ROOT / "SkillCert_Chapters_3_to_5.docx"
MERGED_OUTPUT = ROOT / "SkillCert_Complete_Project_Ch1_to_Ch5_Draft.docx"
CHAPTER_ONE_TWO = Path("/home/anzicle/Download/PROJECT chapter one to chapter two.docx")


def is_table_separator(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and set(stripped.replace("|", "").replace(":", "").replace("-", "").strip()) == set()


def parse_table(lines: list[str], start: int):
    table_lines = []
    idx = start
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        table_lines.append(lines[idx].strip())
        idx += 1

    rows = []
    for line in table_lines:
        if is_table_separator(line):
            continue
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        rows.append(cells)
    return rows, idx


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = document.add_table(rows=len(rows), cols=max(len(row) for row in rows))
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = value
            if r_idx == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def add_heading_safe(document: Document, text: str, level: int):
    try:
        return document.add_heading(text, level=level)
    except KeyError:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(text)
        run.bold = True
        if level == 1:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run.font.size = Pt(14)
        else:
            run.font.size = Pt(12)
        return paragraph


def add_list_item(document: Document, line: str) -> None:
    text = line.strip()
    if text[:2].isdigit() and ". " in text[:4]:
        item = text.split(". ", 1)[1]
        try:
            document.add_paragraph(item, style="List Number")
        except KeyError:
            document.add_paragraph(text)
    else:
        try:
            document.add_paragraph(text[2:].strip(), style="List Bullet")
        except KeyError:
            document.add_paragraph(text)


def append_markdown(document: Document) -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    idx = 0
    while idx < len(lines):
        line = lines[idx].rstrip()
        stripped = line.strip()

        if not stripped:
            idx += 1
            continue

        if stripped.startswith("|"):
            rows, idx = parse_table(lines, idx)
            add_table(document, rows)
            continue

        if stripped.startswith("# "):
            paragraph = add_heading_safe(document, stripped[2:].strip(), level=1)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif stripped.startswith("## "):
            add_heading_safe(document, stripped[3:].strip(), level=2)
        elif stripped.startswith("### "):
            add_heading_safe(document, stripped[4:].strip(), level=3)
        elif stripped.startswith("#### "):
            add_heading_safe(document, stripped[5:].strip(), level=4)
        elif stripped.startswith("- ") or (stripped[:2].isdigit() and ". " in stripped[:4]):
            add_list_item(document, stripped)
        elif stripped.startswith("**Figure"):
            paragraph = document.add_paragraph(stripped.replace("**", ""))
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.italic = True
        else:
            document.add_paragraph(stripped)

        idx += 1


def apply_page_settings(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)


def build_docx() -> None:
    document = Document()
    apply_page_settings(document)
    append_markdown(document)

    document.save(OUTPUT)
    print(OUTPUT)

    if CHAPTER_ONE_TWO.exists():
        merged = Document(CHAPTER_ONE_TWO)
        apply_page_settings(merged)
        merged.add_page_break()
        append_markdown(merged)
        merged.save(MERGED_OUTPUT)
        print(MERGED_OUTPUT)


if __name__ == "__main__":
    build_docx()
