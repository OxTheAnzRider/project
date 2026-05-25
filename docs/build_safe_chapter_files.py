from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "chapters_3_to_5.md"
SAFE_DOCX = ROOT / "SkillCert_Chapters_3_to_5_SAFE.docx"
SAFE_RTF = ROOT / "SkillCert_Chapters_3_to_5_SAFE.rtf"
SAFE_TXT = ROOT / "SkillCert_Chapters_3_to_5_SAFE.txt"


def clean_markdown_line(line: str) -> str:
    text = line.strip()
    if text.startswith("#"):
        text = text.lstrip("#").strip()
    text = text.replace("**", "")
    text = text.replace("*", "")
    text = text.replace("`", "")
    if text.startswith("|"):
        cells = [cell.strip() for cell in text.strip("|").split("|")]
        if all(set(cell.replace(":", "").replace("-", "").strip()) <= set() for cell in cells):
            return ""
        text = " | ".join(cells)
    return text


def build_safe_docx() -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    for raw in SOURCE.read_text(encoding="utf-8").splitlines():
        if not raw.strip():
            continue
        text = clean_markdown_line(raw)
        if not text:
            continue
        paragraph = document.add_paragraph()
        run = paragraph.add_run(text)
        if raw.startswith("# "):
            run.bold = True
            run.font.size = Pt(14)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif raw.startswith("## ") or raw.startswith("### ") or raw.startswith("#### "):
            run.bold = True
            run.font.size = Pt(12)

    document.save(SAFE_DOCX)


def rtf_escape(text: str) -> str:
    return (
        text.replace("\\", "\\\\")
        .replace("{", "\\{")
        .replace("}", "\\}")
        .replace("\n", "\\par\n")
    )


def build_safe_rtf_and_txt() -> None:
    plain_lines = []
    rtf_lines = [r"{\rtf1\ansi\deff0", r"{\fonttbl{\f0 Times New Roman;}}", r"\fs24"]
    for raw in SOURCE.read_text(encoding="utf-8").splitlines():
        text = clean_markdown_line(raw)
        if not text:
            plain_lines.append("")
            rtf_lines.append(r"\par")
            continue
        plain_lines.append(text)
        if raw.startswith("# "):
            rtf_lines.append(r"\b\fs28 " + rtf_escape(text) + r"\b0\fs24\par")
        elif raw.startswith("## ") or raw.startswith("### ") or raw.startswith("#### "):
            rtf_lines.append(r"\b " + rtf_escape(text) + r"\b0\par")
        else:
            rtf_lines.append(rtf_escape(text) + r"\par")
    rtf_lines.append("}")
    SAFE_TXT.write_text("\n".join(plain_lines), encoding="utf-8")
    SAFE_RTF.write_text("\n".join(rtf_lines), encoding="utf-8")


def main() -> None:
    build_safe_docx()
    build_safe_rtf_and_txt()
    print(SAFE_DOCX)
    print(SAFE_RTF)
    print(SAFE_TXT)


if __name__ == "__main__":
    main()
